import streamlit as st
from docx import Document
import os
import random
import re

# Path to the folder containing question papers
QUESTION_PAPER_DIR = "./docs"

# Function to parse the .docx file and get all the questions
def get_questions_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

        # Regex pattern to match anything starting with ***<number> and ending with ".***" or " ***"
        pattern = re.compile(r"\*{3}\d+.*?(?:\. \*{3}|\.{3}|\s\*{3})", re.DOTALL)
        questions = pattern.findall(full_text)

        return [q.strip() for q in questions]
    except Exception as e:
        st.error(f"Failed to read file {file_path}: {e}")
        return []

# Load available question paper files from the docs folder
def load_question_papers():
    if not os.path.exists(QUESTION_PAPER_DIR):
        os.makedirs(QUESTION_PAPER_DIR)
    files = [f for f in os.listdir(QUESTION_PAPER_DIR) if f.endswith(".docx")]
    return files

# Combine selected questions and create a new document
def generate_combined_document(selected_questions):
    doc = Document()
    for exam_code, questions in selected_questions:
        doc.add_heading(f"Questions from {exam_code}", level=1)
        for question in questions:
            doc.add_paragraph(question)
    doc_path = "combined_questions.docx"
    doc.save(doc_path)
    return doc_path

# Streamlit UI
def main():
    st.title("Question Setter")

    if "selected_files" not in st.session_state:
        st.session_state.selected_files = []

    files = load_question_papers()
    exam_codes = [f.replace(".docx", "") for f in files]
    st.sidebar.header("Search and Select Exam Code")

    search_code = st.sidebar.text_input("Enter Exam Code:")

    if search_code:
        matching_codes = [code for code in exam_codes if search_code.lower() in code.lower()]
    else:
        matching_codes = exam_codes

    if matching_codes:
        exam_code = st.selectbox(
            "Select Exam Code", 
            matching_codes, 
            key=f"exam_code_{len(st.session_state.selected_files)}"
        )

        if exam_code:
            questions = get_questions_from_docx(f"{QUESTION_PAPER_DIR}/{exam_code}.docx")

            if questions:
                st.write(f"Exam Code: {exam_code}")
                st.write(f"Total Questions in the file: {len(questions)}")

                num_questions = st.number_input(
                    "How many questions to select?",
                    min_value=1,
                    max_value=len(questions),
                    step=1,
                    key=f"num_questions_{len(st.session_state.selected_files)}"
                )

                if st.button(f"Add Questions from {exam_code}"):
                    selected = random.sample(questions, num_questions)
                    st.session_state.selected_files.append({
                        "exam_code": exam_code,
                        "selected_questions": selected,
                        "num_questions": num_questions
                    })
                    st.success(f"Selected {num_questions} questions from {exam_code}")

    if st.session_state.selected_files:
        st.subheader("Selected Exam Codes and Number of Questions")

        st.write("No. | Exam Code | Number of Questions")
        for idx, selected_file in enumerate(st.session_state.selected_files, start=1):
            st.write(f"{idx} | {selected_file['exam_code']} | {selected_file['num_questions']}")

    if st.button("Proceed to next step"):
        if not st.session_state.selected_files:
            st.warning("Please select at least one set of questions.")
        else:
            combined_doc_path = generate_combined_document(
                [(file['exam_code'], file['selected_questions']) for file in st.session_state.selected_files]
            )
            st.success("Questions Combined Successfully!")
            st.download_button(
                label="Download Combined Questions",
                data=open(combined_doc_path, "rb").read(),
                file_name="combined_questions.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# Correct entry point
if __name__ == "__main__":
    main()
